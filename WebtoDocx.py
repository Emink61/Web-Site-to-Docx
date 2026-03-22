from bs4 import BeautifulSoup as bs
import requests
from docx import Document
import os

documentName = input('Enter Document name:')
fileName = input('Enter File name:')
url = input('Enter your URL:')
os.makedirs(documentName, exist_ok=True)
try:
    response = requests.get(url, timeout=10)
    soup = bs(response.text, 'html.parser')

    doc = Document()
    doc.add_heading(fileName, 0)

    ContentBox = soup.find('div', class_ = 'enter the div class you want to read')
    if ContentBox:
        for cont in ContentBox.find_all(['h1', 'h2', 'h3', 'p', 'li']):
            content = cont.get_text().strip()
            if not content:
                continue
            if cont.name == 'h1':
                doc.add_heading(content, level=1)
            elif cont.name == 'h2':
                doc.add_heading(content, level=2)
            elif cont.name == 'h3':
                doc.add_heading(content, level=3)
            elif cont.name == 'li':
                doc.add_paragraph(content, style='List Bullet')
            else:
                doc.add_paragraph(content)
        doc.save(f"{documentName}/{fileName}.docx")

    else:
        print("Cant Get It!")
except Exception as e:
    print(e)